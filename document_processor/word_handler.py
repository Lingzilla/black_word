import re
from typing import List, Tuple
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Pt

class WordHandler:
    def __init__(self):
        # 定义敏感信息的正则表达式模式
        self.patterns = [
            r'\d{17}[\dXx]',  # 身份证号
            r'\d{11}',        # 手机号
            r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',  # 邮箱
            # 可以添加更多模式
        ]

    def process_document(self, input_path: str, output_path: str) -> None:
        """处理Word文档，标记并涂黑敏感信息"""
        doc = Document(input_path)
        
        # 处理每个段落
        for paragraph in doc.paragraphs:
            self._process_paragraph(paragraph)
            
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_paragraph(paragraph)
        
        # 保存修改后的文档
        doc.save(output_path)

    def _process_paragraph(self, paragraph) -> None:
        """处理单个段落中的敏感信息"""
        if not paragraph.text.strip():
            return

        # 存储原始文本和runs
        original_text = paragraph.text
        sensitive_ranges = self._find_sensitive_content(original_text)

        if not sensitive_ranges:
            return

        # 清空段落
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

        # 重新构建段落，给敏感信息涂黑
        current_pos = 0
        for start, end in sensitive_ranges:
            # 添加普通文本
            if current_pos < start:
                run = paragraph.add_run(original_text[current_pos:start])
                run.font.size = Pt(12)

            # 添加涂黑的敏感文本
            sensitive_text = original_text[start:end]
            run = paragraph.add_run(sensitive_text)
            run.font.highlight_color = RGBColor(0, 0, 0)  # 黑色高亮
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色文字
            run.font.size = Pt(12)

            current_pos = end

        # 添加剩余的普通文本
        if current_pos < len(original_text):
            run = paragraph.add_run(original_text[current_pos:])
            run.font.size = Pt(12)

    def _find_sensitive_content(self, text: str) -> List[Tuple[int, int]]:
        """查找文本中的敏感内容，返回位置范围列表"""
        ranges = []
        for pattern in self.patterns:
            for match in re.finditer(pattern, text):
                ranges.append((match.start(), match.end()))
        
        # 合并重叠区域
        if ranges:
            ranges.sort(key=lambda x: x[0])
            merged = [ranges[0]]
            for current in ranges[1:]:
                if current[0] <= merged[-1][1]:
                    merged[-1] = (merged[-1][0], max(merged[-1][1], current[1]))
                else:
                    merged.append(current)
            return merged
        return []